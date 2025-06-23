import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO

st.set_page_config(page_title="MS WORD Format Checker")

st.title("🧾 MS WORD Format Checker")
st.write("कृपया नीचे विवरण भरें और .docx फाइल अपलोड करें:")

name = st.text_input("👤 छात्र का नाम")
roll = st.text_input("🆔 रोल नंबर")
uploaded_file = st.file_uploader("अपनी Word फ़ाइल अपलोड करें (.docx)", type=["docx"])

def check_formatting(file):
    doc = Document(file)
    score = 0
    feedback = []

    section = doc.sections[0]

    # Q1(A): Page size A3
    width = round(section.page_width.inches, 1)
    height = round(section.page_height.inches, 1)
    if width == 11.7 and height == 16.5:
        score += 2
        feedback.append("✅ Q1(A): पेज साइज A3 सेट है")
    else:
        feedback.append("❌ Q1(A): पेज साइज A3 नहीं है")

    # Q1(B): Paragraph border in 2nd paragraph
    if len(doc.paragraphs) > 1 and 'w:top' in doc.paragraphs[1]._element.xml:
        score += 2
        feedback.append("✅ Q1(B): दूसरे पैराग्राफ में बॉर्डर है")
    else:
        feedback.append("❌ Q1(B): दूसरे पैराग्राफ में बॉर्डर नहीं है")

    # Q1(C): Top margin 0.6
    top_margin = round(section.top_margin.inches, 1)
    if top_margin == 0.6:
        score += 2
        feedback.append("✅ Q1(C): टॉप मार्जिन 0.6 है")
    else:
        feedback.append("❌ Q1(C): टॉप मार्जिन 0.6 नहीं है")

    # Q1(D): 'Define' bold and highlighted in para 4
    if len(doc.paragraphs) > 3:
        found = any("Define" in run.text and run.bold and run.font.highlight_color for run in doc.paragraphs[3].runs)
        if found:
            score += 4
            feedback.append("✅ Q1(D): 'Define' बोल्ड और हाईलाइट किया गया है")
        else:
            feedback.append("❌ Q1(D): 'Define' बोल्ड और हाईलाइट नहीं किया गया")

    return score, feedback

if uploaded_file and name and roll:
    score, results = check_formatting(BytesIO(uploaded_file.read()))
    st.success(f"🎯 {name} (Roll: {roll}) – Total Score: {score}/50")
    st.write("### 📋 फीडबैक:")
    for r in results:
        st.write(r)

    # Downloadable report
    report = f"नाम: {name}\nरोल: {roll}\nअंक: {score}/50\n\n" + "\n".join(results)
    st.download_button("📥 रिपोर्ट डाउनलोड करें", report, file_name=f"{roll}_report.txt", mime="text/plain")
elif uploaded_file and (not name or not roll):
    st.warning("कृपया नाम और रोल नंबर भरें")
